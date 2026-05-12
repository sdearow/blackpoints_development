"""Genera la relazione metodologica in formato Word (.docx).

Versione professionale con immagini integrate dalla presentazione PPT
e formattazione avanzata (header/footer, copertina, didascalie figure).
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from pathlib import Path

IMG_DIR = Path(__file__).resolve().parent / "pptx_images"

# Colori corporate RSM / Roma
_BORDEAUX = RGBColor(0x6B, 0x11, 0x27)
_BLU_SCURO = RGBColor(0x0D, 0x2B, 0x52)
_GRIGIO = RGBColor(0x33, 0x33, 0x33)
_GRIGIO_CHIARO = RGBColor(0x66, 0x66, 0x66)


def _stile_doc(doc: Document) -> None:
    """Configura stili base del documento."""
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    for lvl, (nome, sz, colore) in enumerate([
        ("Heading 1", 18, _BLU_SCURO),
        ("Heading 2", 14, _BLU_SCURO),
        ("Heading 3", 12, _BORDEAUX),
    ]):
        s = doc.styles[nome]
        s.font.name = "Calibri"
        s.font.size = Pt(sz)
        s.font.bold = True
        s.font.color.rgb = colore
        s.paragraph_format.space_before = Pt(18 if lvl == 0 else 14)
        s.paragraph_format.space_after = Pt(6)


def _p(doc, testo, style="Normal", bold=False, italic=False, align=None):
    """Aggiunge un paragrafo con formattazione opzionale."""
    p = doc.add_paragraph(testo, style=style)
    if bold or italic:
        for run in p.runs:
            run.bold = bold
            run.italic = italic
    if align:
        p.alignment = align
    return p


def _figura(doc, img_path, caption, width=Cm(15)):
    """Inserisce un'immagine centrata con didascalia numerata."""
    if not img_path.exists():
        _p(doc, f"[Immagine non trovata: {img_path.name}]", italic=True)
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=width)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(caption)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = _GRIGIO_CHIARO


def _linea_separatore(doc):
    """Inserisce una linea orizzontale sottile."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="999999"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def _pagina_copertina(doc):
    """Genera la pagina di copertina professionale."""
    for _ in range(4):
        doc.add_paragraph()

    logo_path = IMG_DIR / "slide01_img3_Logo_500x500_png.png"
    if logo_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(logo_path), width=Cm(5))

    doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("RELAZIONE METODOLOGICA")
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = _BLU_SCURO

    doc.add_paragraph()

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = st.add_run(
        "Sistema di Identificazione dei Black Point Incidentali\n"
        "Rete Stradale del Comune di Roma Capitale"
    )
    r2.font.size = Pt(16)
    r2.font.color.rgb = _BORDEAUX

    for _ in range(4):
        doc.add_paragraph()

    _linea_separatore(doc)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.paragraph_format.space_before = Pt(12)
    for label, value in [
        ("Committente: ", "Roma Servizi per la Mobilità S.r.l."),
        ("Metodologia: ", "Empirical Bayes – Highway Safety Manual (AASHTO)"),
        ("Periodo di analisi: ", "2019 – 2024"),
    ]:
        r_l = info.add_run(label)
        r_l.font.size = Pt(10)
        r_l.font.bold = True
        r_l.font.color.rgb = _GRIGIO
        r_v = info.add_run(value + "\n")
        r_v.font.size = Pt(10)
        r_v.font.color.rgb = _GRIGIO

    doc.add_page_break()


def _indice(doc):
    """Inserisce un sommario manuale."""
    doc.add_heading("Indice", level=1)

    voci = [
        ("1.", "Premessa e obiettivi del lavoro"),
        ("2.", "Architettura del sistema"),
        ("3.", "Dati di input"),
        ("4.", "Fase 0 – Acquisizione e pulizia dei dati"),
        ("5.", "Fase 1 – Matching spaziale incidenti-rete"),
        ("6.", "Fase 2 – Safety Performance Functions (SPF)"),
        ("7.", "Fase 3 – Stima Empirical Bayes e EPDO"),
        ("8.", "Fase 4 – Indice Composito di Priorità (ICP)"),
        ("9.", "Fase 5 – Visualizzazione e reporting decisionale"),
        ("10.", "Limiti del prototipo e sviluppi futuri"),
        ("11.", "Riferimenti bibliografici"),
    ]

    for num, titolo in voci:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r_n = p.add_run(f"{num}  ")
        r_n.font.bold = True
        r_n.font.size = Pt(11)
        r_n.font.color.rgb = _BLU_SCURO
        r_t = p.add_run(titolo)
        r_t.font.size = Pt(11)

    doc.add_page_break()


def genera():
    doc = Document()
    _stile_doc(doc)

    # ================================================================
    # COPERTINA
    # ================================================================
    _pagina_copertina(doc)

    # ================================================================
    # INDICE
    # ================================================================
    _indice(doc)

    # ================================================================
    # 1. PREMESSA
    # ================================================================
    doc.add_heading("1. Premessa e obiettivi del lavoro", level=1)

    _p(doc, (
        "Il presente documento descrive la metodologia adottata per la costruzione "
        "di un sistema di identificazione e classificazione dei black point "
        "incidentali sulla rete stradale del Comune di Roma Capitale."
    ))

    _p(doc, (
        "Il sistema nasce dall'esigenza di dotare l'Amministrazione di uno strumento "
        "oggettivo e scientificamente fondato per l'individuazione dei siti della rete "
        "stradale che presentano una concentrazione anomala di incidenti, al fine di "
        "orientare la programmazione degli interventi di messa in sicurezza verso le "
        "criticità effettive e misurabili. L'approccio supera i limiti dei metodi "
        "tradizionali basati sul semplice conteggio degli incidenti, che sono soggetti "
        "all'effetto regression-to-the-mean e non distinguono tra siti strutturalmente "
        "pericolosi e siti dove la concentrazione di eventi è casuale."
    ))

    _p(doc, (
        "La metodologia si fonda sui principi della Network Safety Analysis descritti "
        "nell'Highway Safety Manual (AASHTO, 2010) e nella letteratura internazionale "
        "sulla sicurezza stradale, con particolare riferimento al metodo Empirical Bayes "
        "(Hauer, 1997) e alla normativa europea sulla gestione della sicurezza delle "
        "infrastrutture stradali (Direttiva 2008/96/CE, aggiornata dalla Direttiva "
        "2019/1936). Il sistema è coerente con i paradigmi Vision Zero e Safe System, "
        "che pongono al centro l'inaccettabilità delle morti e dei feriti gravi sulla strada "
        "e la necessità di un approccio proattivo alla sicurezza."
    ))

    _p(doc, (
        "L'obiettivo finale è duplice: da un lato, fornire una classifica dei siti "
        "prioritari basata su un indice composito multi-criterio che integri frequenza, "
        "gravità, vulnerabilità degli utenti deboli e condizioni operative della strada; "
        "dall'altro, mettere a disposizione una dashboard interattiva per l'esplorazione "
        "dei risultati, l'analisi di sensitività e il supporto alle decisioni."
    ))

    # ================================================================
    # 2. ARCHITETTURA DEL SISTEMA
    # ================================================================
    doc.add_heading("2. Architettura del sistema", level=1)

    _p(doc, (
        "Il sistema è strutturato come una pipeline di elaborazione sequenziale "
        "composta da sei fasi, ciascuna delle quali produce output intermedi "
        "utilizzati dalle fasi successive. L'intera pipeline è implementata in Python "
        "ed è completamente riproducibile a partire dai dati grezzi: ogni esecuzione "
        "con gli stessi dati di input e gli stessi parametri di configurazione produce "
        "risultati identici."
    ))

    _p(doc, (
        "Il diagramma seguente sintetizza le sei fasi della pipeline e il loro "
        "contenuto principale."
    ))

    _figura(doc,
            IMG_DIR / "slide04_img1_Immagine_2.png",
            "Figura 1 – Architettura della pipeline: le sei fasi del sistema Black Point.",
            width=Cm(16))

    _p(doc, (
        "Il sistema di riferimento metrico utilizzato per tutti i calcoli geometrici "
        "è UTM zona 33N (EPSG:32633); la visualizzazione finale avviene in WGS84 "
        "(EPSG:4326). Tutti i parametri operativi (soglie, pesi, percorsi dei file) "
        "sono centralizzati in un file di configurazione YAML, modificabile senza "
        "intervenire sul codice."
    ))

    # ================================================================
    # 3. DATI DI INPUT
    # ================================================================
    doc.add_heading("3. Dati di input", level=1)

    _p(doc, (
        "Il sistema integra quattro fonti dati principali, ciascuna delle quali "
        "contribuisce informazioni essenziali per la costruzione del modello."
    ))

    doc.add_heading("3.1. Database incidentale del Comune di Roma", level=2)
    _p(doc, (
        "Il database incidentale è costituito da più file CSV che coprono il periodo "
        "2004–2024, derivanti da diverse campagne di rilevazione e rigeolocalizzazione. "
        "Comprende il dataset storico (2004–2022), estratti intermedi e le "
        "rigeolocalizzazioni annuali (2022, 2023, 2024) che hanno priorità in caso di "
        "duplicati per lo stesso incidente. Per il periodo di analisi (2019–2024) sono "
        "disponibili circa 90.000 record."
    ))

    doc.add_heading("3.2. Grafo stradale TomTom", level=2)
    _p(doc, (
        "La rete stradale di riferimento è il grafo TomTom 2024, composto da circa "
        "94.000 archi con geometria MultiLineString e attributi di traffico. Per ogni arco "
        "sono disponibili il Traffico Giornaliero Medio (TGM), la distribuzione completa "
        "delle velocità istantanee rilevate dai probe GPS (percentili dal 5° al 95°, in "
        "particolare la V85 come proxy della velocità operativa), il limite di velocità "
        "vigente e la classe funzionale FRC (scala TomTom 0–7)."
    ))

    doc.add_heading("3.3. Grafo PGTU 2026 del Comune", level=2)
    _p(doc, (
        "Il Piano Generale del Traffico Urbano 2026 fornisce la classificazione funzionale "
        "aggiornata delle strade comunali (Scorrimento, Interquartiere, Interzonale, "
        "Quartiere), i flag di grande viabilità e di trasporto pubblico locale. Questa "
        "classificazione viene ereditata dal grafo TomTom tramite join spaziale."
    ))

    doc.add_heading("3.4. Catasto degli impianti semaforici", level=2)
    _p(doc, (
        "Il catasto contiene la posizione geografica e il tipo di ogni impianto semaforico "
        "(veicolare o pedonale). I semafori veicolari vengono utilizzati per definire le "
        "intersezioni semaforizzate, informazione rilevante per la stratificazione dei "
        "modelli SPF."
    ))

    # ================================================================
    # 4. FASE 0
    # ================================================================
    doc.add_heading("4. Fase 0 – Acquisizione e pulizia dei dati", level=1)

    _p(doc, (
        "La prima fase della pipeline ha l'obiettivo di acquisire, standardizzare e "
        "validare i dati di input, producendo dataset puliti e coerenti pronti per "
        "le fasi successive."
    ))

    doc.add_heading("4.1. Pulizia del database incidentale", level=2)
    _p(doc, (
        "La procedura di pulizia esegue le seguenti operazioni principali:"
    ))

    _p(doc, (
        "Deduplica su identificativo univoco: quando lo stesso incidente compare "
        "in più file sorgente, viene conservata la versione proveniente dal file con "
        "priorità più alta (le rigeolocalizzazioni annuali hanno priorità massima "
        "rispetto al dataset storico)."
    ))

    _p(doc, (
        "Standardizzazione delle coordinate: le coordinate originali in Gauss-Boaga "
        "zona 2 (EPSG:3004) vengono riproiettate nel sistema metrico di lavoro "
        "(EPSG:32633). Normalizzazione toponomastica: espansione delle abbreviazioni "
        "tipiche della viabilità romana (V. → Via, P.zza → Piazza, V.le → Viale) e "
        "uniformazione delle maiuscole."
    ))

    _p(doc, (
        "Classificazione della gravità in tre livelli: mortale (almeno un decesso), "
        "con feriti (almeno un ferito, nessun decesso), solo danni materiali. Il dataset "
        "originale non distingue tra feriti gravi e lievi; si adotta pertanto un'unica "
        "categoria «feriti»."
    ))

    _p(doc, (
        "Calcolo di un flag di qualità della geocodifica (alta, media, bassa) basato "
        "sulla presenza di coordinate valide, sul campo di conferma e sulla specificità "
        "della localizzazione. Filtro spaziale entro il confine comunale e filtro di "
        "qualità: per la calibrazione dei modelli vengono utilizzati solo gli incidenti "
        "con geocodifica di qualità «alta»."
    ))

    doc.add_heading("4.2. Preparazione della rete stradale", level=2)
    _p(doc, (
        "La preparazione della rete include la validazione topologica (rimozione degli "
        "archi con geometria nulla o invalida), il calcolo delle covariate derivate per "
        "i modelli SPF — log(TGM), log(lunghezza in km), IQR normalizzato delle velocità, "
        "rapporto V85/limite — e il join spaziale con il grafo PGTU per ereditare la "
        "classificazione funzionale aggiornata."
    ))

    _p(doc, (
        "Il catasto dei semafori viene normalizzato con le stesse regole toponomastiche "
        "del database incidentale e corretto per un offset sistematico calcolato "
        "empiricamente. Solo i semafori veicolari vengono utilizzati per la definizione "
        "delle intersezioni semaforizzate."
    ))

    # ================================================================
    # 5. FASE 1
    # ================================================================
    doc.add_heading("5. Fase 1 – Matching spaziale incidenti-rete", level=1)

    _p(doc, (
        "La Fase 1 ha il compito di costruire i «siti» della rete (intersezioni e "
        "segmenti omogenei) e di assegnare ciascun incidente al sito di competenza. "
        "Questa operazione è fondamentale perché i modelli SPF e l'Empirical Bayes "
        "lavorano a livello di sito, non di singolo evento."
    ))

    doc.add_heading("5.1. Estrazione delle intersezioni", level=2)
    _p(doc, (
        "Le intersezioni vengono estratte automaticamente dal grafo TomTom identificando "
        "i nodi con grado topologico ≥ 3 (almeno tre archi convergenti). L'algoritmo "
        "utilizza un clustering degli endpoint entro 1.5 m tramite KD-tree e union-find "
        "per gestire i disallineamenti topologici."
    ))

    _p(doc, (
        "Il grafo TomTom genera un numero elevato di falsi positivi: nodi a grado 3–4 "
        "che non corrispondono a veri incroci ma a confluenze di carreggiate separate, "
        "svincoli a livelli sfalsati o micro-segmentazioni della rete. Per filtrare questi "
        "artefatti si applicano tre criteri progressivi: filtro mono-toponimo, filtro FRC "
        "uniforme e clustering di prossimità (nodi entro 30 m con toponimo condiviso "
        "vengono fusi). Dopo il filtraggio, ogni intersezione viene arricchita con "
        "l'informazione di semaforizzazione tramite join spaziale."
    ))

    doc.add_heading("5.2. Costruzione dei segmenti omogenei", level=2)
    _p(doc, (
        "I segmenti omogenei vengono costruiti concatenando archi TomTom consecutivi "
        "che condividono lo stesso toponimo e presentano una variazione del TGM contenuta "
        "(< 30%). La segmentazione si interrompe a ogni intersezione reale, al cambio di "
        "toponimo, quando la variazione relativa del TGM supera la soglia, o quando la "
        "lunghezza cumulata supera 2.000 m."
    ))

    _p(doc, (
        "Per ciascun segmento vengono calcolati gli attributi aggregati: TGM medio, V85 media, "
        "limite di velocità medio, IQR normalizzato medio (tutti pesati per lunghezza dell'arco), "
        "classe FRC modale e classificazione PGTU modale."
    ))

    doc.add_heading("5.3. Assegnazione degli incidenti alla rete", level=2)
    _p(doc, (
        "L'abbinamento avviene secondo una gerarchia di criteri con priorità decrescente:"
    ))

    _p(doc, (
        "1) Intersezione: se l'incidente ricade entro un raggio di 25 m da un "
        "nodo-intersezione, viene assegnato a quell'intersezione. Questa priorità "
        "riflette le caratteristiche specifiche degli incidenti agli incroci (conflitti "
        "di traiettoria, svolta, attraversamento pedonale).\n"
        "2) Segmento geometrico: assegnazione al segmento più vicino entro 30 m.\n"
        "3) Fallback toponomastico: ricerca di un segmento entro 100 m con toponimo "
        "compatibile (matching fuzzy con soglia ≥ 85).\n"
        "4) Non abbinato: gli incidenti residui restano esclusi dall'analisi."
    ))

    # ================================================================
    # 6. FASE 2 – SPF
    # ================================================================
    doc.add_heading("6. Fase 2 – Safety Performance Functions (SPF)", level=1)

    _p(doc, (
        "Le Safety Performance Functions sono modelli di regressione che stimano il "
        "numero atteso di incidenti per un sito sulla base delle sue caratteristiche "
        "di traffico e geometria. Il valore predetto dalla SPF rappresenta la performance "
        "«media» della rete per quella classe di siti, ed è il riferimento rispetto al "
        "quale l'Empirical Bayes calcola l'eccesso o il deficit di incidentalità."
    ))

    doc.add_heading("6.1. Struttura del modello", level=2)
    _p(doc, (
        "Il modello adottato è la regressione binomiale negativa di tipo NB2 (Cameron e "
        "Trivedi, 1998), standard nella letteratura sulla sicurezza stradale per la sua "
        "capacità di gestire la sovradispersione tipica dei conteggi incidentali. La "
        "distribuzione binomiale negativa introduce un parametro aggiuntivo α (overdispersion) "
        "rispetto alla distribuzione di Poisson; il suo inverso k = 1/α è fondamentale per il "
        "calcolo del peso Empirical Bayes."
    ))

    _p(doc, (
        "La figura seguente mostra la struttura formale dei modelli SPF per segmenti e "
        "intersezioni."
    ))

    _figura(doc,
            IMG_DIR / "slide07_img1_Immagine_4.png",
            "Figura 2 – Formulazione dei modelli SPF per segmenti stradali e intersezioni.",
            width=Cm(14))

    _p(doc, (
        "Per i segmenti, il modello base include il logaritmo del TGM e della lunghezza "
        "come covariate principali, con un offset per il numero di anni di osservazione. "
        "Il modello esteso aggiunge la V85 (velocità operativa) e l'IQR normalizzato "
        "(dispersione delle velocità). Per le intersezioni, la covariata principale è il "
        "flusso entrante (somma dei TGM degli archi convergenti divisa per 2), con "
        "l'eventuale aggiunta del numero di bracci nel modello esteso."
    ))

    _p(doc, (
        "Il modello esteso viene adottato solo se produce un miglioramento dell'AIC "
        "(Akaike Information Criterion) rispetto al modello base, evitando il rischio "
        "di overfitting."
    ))

    doc.add_heading("6.2. Stratificazione per categoria", level=2)
    _p(doc, (
        "I modelli SPF vengono calibrati separatamente per ciascuna categoria funzionale, "
        "poiché le relazioni tra traffico e incidentalità differiscono significativamente "
        "tra tipologie di strada. Per i segmenti la categoria SPF è derivata dalla "
        "classificazione PGTU 2026 (IQ – Interquartiere, IZ – Interzonale, Q – Quartiere, "
        "LOCALE, EXTRAURBANA). Per le intersezioni la stratificazione avviene per stato di "
        "semaforizzazione (semaforizzata / non semaforizzata). Le categorie con meno di "
        "50 siti vengono accorpate."
    ))

    # ================================================================
    # 7. FASE 3 – EB e EPDO
    # ================================================================
    doc.add_heading("7. Fase 3 – Stima Empirical Bayes e EPDO", level=1)

    _p(doc, (
        "Il metodo Empirical Bayes (Hauer, 1997) rappresenta lo stato dell'arte per la "
        "stima della sicurezza di un sito stradale. È il metodo raccomandato dall'Highway "
        "Safety Manual e adottato dalle principali agenzie stradali internazionali. Il suo "
        "vantaggio fondamentale è la capacità di correggere l'effetto regression-to-the-mean, "
        "che nei metodi tradizionali porta a sovrastimare la pericolosità di siti dove si è "
        "verificata una concentrazione casuale di eventi."
    ))

    doc.add_heading("7.1. Il metodo Empirical Bayes", level=2)

    _p(doc, (
        "L'EB combina due fonti di informazione complementari: la predizione del modello "
        "SPF (E), che rappresenta la performance media di siti strutturalmente simili, e il "
        "conteggio osservato (O), che riflette le condizioni reali ma è soggetto a "
        "fluttuazione casuale."
    ))

    _figura(doc,
            IMG_DIR / "slide09_img1_Immagine_2.png",
            "Figura 3 – Schema del metodo Empirical Bayes e della valutazione della severità.",
            width=Cm(16))

    _p(doc, (
        "La stima EB combina le due fonti assegnando un peso di affidabilità (w) che dipende "
        "dalla precisione relativa del modello rispetto all'osservato:"
    ))

    _p(doc, (
        "w = 1 / (1 + E · k)        dove k = 1/α (parametro di sovradispersione NB2)\n"
        "EB = w · E + (1 − w) · O\n"
        "Excess = EB − E"
    ), italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    _p(doc, (
        "Quando E è grande (siti ad alto flusso con molti dati), il peso w tende a zero e "
        "la stima EB coincide con l'osservato. Quando E è piccolo (siti a basso flusso), w "
        "tende a 1 e la stima converge verso la predizione SPF, proteggendo dall'effetto "
        "regression-to-the-mean. L'eccesso atteso (Excess = EB − E) quantifica quanti "
        "incidenti in più il sito ha rispetto a un sito medio con le stesse caratteristiche: "
        "un eccesso positivo indica un potenziale black point."
    ))

    doc.add_heading("7.2. EPDO – Equivalent Property Damage Only", level=2)
    _p(doc, (
        "Il conteggio grezzo degli incidenti non distingue tra un tamponamento con soli "
        "danni materiali e un incidente mortale. L'EPDO pondera gli incidenti per gravità "
        "attribuendo pesi relativi secondo la scala classica Hauer/AASHTO:"
    ))

    table = doc.add_table(rows=4, cols=2)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (cat, peso) in enumerate([
        ("Categoria", "Peso EPDO"),
        ("Incidente mortale", "12"),
        ("Incidente con feriti", "3"),
        ("Solo danni materiali", "1"),
    ]):
        table.rows[i].cells[0].text = cat
        table.rows[i].cells[1].text = peso
        if i == 0:
            for cell in table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_paragraph()

    _p(doc, (
        "I pesi EPDO 12/3/1 sono moderati e consolidati nella letteratura, evitando che "
        "un singolo evento mortale casuale domini la classifica. L'eccesso EB viene pesato "
        "con il rapporto EPDO medio del sito per ottenere l'excess_EPDO, che costituisce "
        "la componente A dell'indice composito."
    ))

    doc.add_heading("7.3. Costo sociale", level=2)
    _p(doc, (
        "Parallelamente all'EPDO, il sistema calcola il costo sociale dell'eccesso di "
        "incidentalità utilizzando i costi unitari MEF/ISTAT (valori 2022): decesso "
        "1.500.000 €, ferito (blend 15% grave + 85% lieve) 48.300 €, solo danni 9.000 €. "
        "Il costo sociale non entra nel calcolo dell'ICP ma viene riportato nei risultati "
        "per consentire analisi costi-benefici degli interventi infrastrutturali."
    ))

    # ================================================================
    # 8. FASE 4 – ICP
    # ================================================================
    doc.add_heading("8. Fase 4 – Indice Composito di Priorità (ICP)", level=1)

    _p(doc, (
        "L'indice composito di priorità (ICP) rappresenta il cuore del sistema di "
        "classificazione. Aggrega quattro componenti complementari, ciascuna delle "
        "quali cattura un aspetto diverso della pericolosità del sito, in un unico "
        "indicatore sintetico che consente l'ordinamento dei siti per priorità di "
        "intervento."
    ))

    doc.add_heading("8.1. Le quattro componenti dell'ICP", level=2)

    _figura(doc,
            IMG_DIR / "slide10_img1_Immagine_2.png",
            "Figura 4 – Composizione dell'ICP: le quattro componenti e i pesi default.",
            width=Cm(14))

    _p(doc, (
        "Componente A – Eccesso EB pesato (peso default: 40%). È l'excess_EPDO, "
        "ovvero l'eccesso di incidentalità rispetto al modello SPF pesato per la gravità "
        "media del sito. Identifica i siti che hanno statisticamente più incidenti del "
        "previsto, al netto delle fluttuazioni casuali e della composizione per gravità."
    ), bold=True)

    _p(doc, (
        "Componente B – Indice di severità (peso default: 25%). Rapporto tra il "
        "numero di incidenti con esito grave (mortali + feriti) e il totale, moltiplicato "
        "per un fattore di credibilità min(n, 5)/5 che smorza i siti con pochissimi "
        "incidenti dove il rapporto è statisticamente instabile."
    ), bold=True)

    _p(doc, (
        "Componente C – Indice di vulnerabilità utenti deboli (peso default: 20%). "
        "Rapporto tra il numero di incidenti con coinvolgimento di pedoni e il totale, "
        "anch'esso con fattore di credibilità. Penalizza i siti dove la quota di utenti "
        "deboli coinvolti è particolarmente elevata."
    ), bold=True)

    _p(doc, (
        "Componente D – Dispersione delle velocità (peso default: 15%). IQR normalizzato "
        "delle velocità istantanee. Una dispersione elevata segnala situazioni dove coesistono "
        "veicoli molto lenti e molto veloci, tipiche di tratti con geometria confusa o "
        "transizioni tra ambienti stradali diversi."
    ), bold=True)

    doc.add_heading("8.2. Normalizzazione delle componenti", level=2)
    _p(doc, (
        "Le quattro componenti hanno scale e distribuzioni molto diverse. Per renderle "
        "confrontabili, ciascuna viene normalizzata su scala 0–100 tramite percentili "
        "robusti (1° e 99° percentile). Per le componenti con distribuzione fortemente "
        "zero-inflated (tipicamente B e C sui segmenti), si applica una normalizzazione "
        "separata: i valori nulli restano a 0, i valori positivi vengono normalizzati "
        "nell'intervallo [1, 100] sulla sotto-distribuzione dei soli valori > 0."
    ))

    doc.add_heading("8.3. Classificazione in fasce di priorità", level=2)
    _p(doc, (
        "I siti vengono classificati in cinque fasce di priorità basate sui percentili "
        "dell'ICP, come illustrato nella figura seguente."
    ))

    _figura(doc,
            IMG_DIR / "slide11_img1_Immagine_7.png",
            "Figura 5 – Le cinque fasce di priorità e le relative soglie percentili.",
            width=Cm(10))

    doc.add_heading("8.4. Matrice di rischio", level=2)
    _p(doc, (
        "Oltre alla classifica ICP, i siti vengono posizionati in una matrice di rischio "
        "2×2 che incrocia l'eccesso di incidentalità (componente A) con la severità "
        "(componente B). Le soglie sono fissate al 75° percentile calcolato esclusivamente "
        "sui siti con almeno un incidente, per evitare che la massa di siti a zero incidenti "
        "abbassi artificialmente le soglie."
    ))

    _figura(doc,
            IMG_DIR / "slide11_img2_Immagine_9.png",
            "Figura 6 – Matrice di rischio 2×2: i quattro quadranti decisionali.",
            width=Cm(14))

    _p(doc, (
        "I quattro quadranti identificano altrettante strategie di intervento: "
        "Q1 – Intervento urgente (alto eccesso e alta severità), "
        "Q2 – Intervento programmato (alto eccesso, severità contenuta), "
        "Q3 – Indagine approfondita (eccesso contenuto, alta severità), "
        "Q4 – Monitoraggio (entrambi sotto soglia)."
    ))

    # ================================================================
    # 9. FASE 5 – DASHBOARD
    # ================================================================
    doc.add_heading("9. Fase 5 – Visualizzazione e reporting decisionale", level=1)

    _p(doc, (
        "I risultati vengono esportati in diversi formati (GeoJSON, Excel, CSV di sintesi) "
        "e resi accessibili attraverso una dashboard interattiva sviluppata in Dash/Plotly, "
        "che costituisce l'interfaccia principale per l'esplorazione e il supporto alle "
        "decisioni."
    ))

    doc.add_heading("9.1. Mappa interattiva", level=2)
    _p(doc, (
        "La vista principale della dashboard è una mappa interattiva di Roma con i siti "
        "della rete colorati secondo l'intensità dell'ICP (gradiente giallo-arancione-rosso "
        "scuro). L'utente può filtrare per tipo di sito (segmenti/intersezioni), fascia di "
        "priorità e numero di siti da visualizzare (top N per ICP). Selezionando un sito, "
        "il pannello laterale mostra il dettaglio delle quattro componenti tramite un "
        "diagramma radar e la distribuzione degli incidenti per gravità."
    ))

    _figura(doc,
            IMG_DIR / "slide12_img1_Immagine_3.png",
            "Figura 7 – Dashboard: mappa interattiva con dettaglio del sito selezionato.",
            width=Cm(16))

    doc.add_heading("9.2. Diagnostica dei modelli", level=2)
    _p(doc, (
        "La tab di diagnostica SPF consente di verificare la bontà di calibrazione dei "
        "modelli: scatter dei valori osservati vs. predetti (binned), distribuzione dei "
        "residui di Pearson, istogramma della distribuzione degli incidenti per sito e "
        "riepilogo tabellare dei modelli calibrati per ogni categoria (n. siti, n. incidenti, "
        "parametro k, valori medi di E e O). La tab di diagnostica EB mostra le distribuzioni "
        "dell'eccesso, del peso w e la classifica per excess_EPDO."
    ))

    _figura(doc,
            IMG_DIR / "slide13_img1_Immagine_12.png",
            "Figura 8 – Dashboard: diagnostica SPF con scatter O vs E e residui di Pearson.",
            width=Cm(16))

    doc.add_heading("9.3. Analisi di sensitività", level=2)
    _p(doc, (
        "La tab di sensitività permette all'utente di variare i pesi delle quattro componenti "
        "dell'ICP e osservare in tempo reale l'effetto sulla classifica dei siti. La dashboard "
        "mostra la correlazione di Spearman con la classifica default e aggiorna la lista dei "
        "top 20 segmenti e top 20 intersezioni, consentendo di valutare la robustezza dei "
        "risultati rispetto alla scelta dei pesi."
    ))

    _figura(doc,
            IMG_DIR / "slide14_img1_Immagine_11.png",
            "Figura 9 – Dashboard: analisi di sensitività dei pesi e classifica top siti.",
            width=Cm(16))

    doc.add_heading("9.4. Vantaggi del sistema", level=2)
    _p(doc, (
        "Il sistema presenta diversi vantaggi rispetto ai metodi tradizionali di "
        "identificazione dei punti neri:"
    ))

    _p(doc, (
        "Approccio multidimensionale: l'ICP integra frequenza, gravità, vulnerabilità "
        "degli utenti deboli e condizioni operative, superando i limiti dei ranking basati "
        "su un unico indicatore.\n\n"
        "Supporto decisionale: la matrice di rischio e la dashboard interattiva forniscono "
        "strumenti immediati per la programmazione degli interventi, con la possibilità di "
        "esplorare scenari alternativi tramite l'analisi di sensitività.\n\n"
        "Riproducibilità: la pipeline è completamente automatizzata e parametrizzata; ogni "
        "esecuzione con gli stessi dati e parametri produce risultati identici, garantendo "
        "trasparenza e verificabilità.\n\n"
        "Aggiornabilità: al rilascio di nuovi dati incidentali o di traffico, il sistema "
        "può essere rieseguito producendo classifiche aggiornate senza interventi manuali."
    ))

    # ================================================================
    # 10. LIMITI E SVILUPPI
    # ================================================================
    doc.add_heading("10. Limiti del prototipo e sviluppi futuri", level=1)

    _p(doc, (
        "Il sistema qui presentato costituisce un prototipo funzionante che dimostra la "
        "fattibilità dell'approccio e produce risultati già utilizzabili per una prima "
        "individuazione delle criticità della rete. Tuttavia, la qualità dei risultati "
        "finali dipende in misura determinante dalla qualità dei dati di input."
    ))

    doc.add_heading("10.1. Qualità dei dati incidentali: il nodo critico", level=2)
    _p(doc, (
        "La criticità principale riguarda il database incidentale del Comune di Roma. "
        "Allo stato attuale:"
    ))

    _p(doc, (
        "• Il dataset storico copre il periodo 2004–2022 ma la geocodifica degli incidenti "
        "pre-2018 presenta imprecisioni significative e non è stata oggetto di una campagna "
        "di rigeolocalizzazione sistematica.\n"
        "• Le rigeolocalizzazioni annuali (2022–2024) migliorano la precisione ma coprono "
        "solo una parte del periodo. Gli incidenti non rigeolocalizzati dipendono interamente "
        "dal matching spaziale/toponomastico, con un tasso inevitabile di errori.\n"
        "• Non è disponibile una distinzione affidabile tra feriti gravi e feriti lievi: "
        "il sistema tratta tutti i feriti come un'unica categoria, perdendo informazione "
        "preziosa per la valutazione della severità.\n"
        "• Il campo «natura dell'incidente» non è sempre compilato in modo omogeneo, "
        "limitando l'affidabilità della componente C (vulnerabilità pedoni)."
    ))

    _p(doc, (
        "Senza un aggiornamento e una correzione sistematica dei dati incidentali, "
        "non è possibile ottenere una calibrazione accurata dei modelli SPF né una "
        "definizione affidabile delle priorità dei black point. Questo punto è "
        "il prerequisito fondamentale per qualsiasi sviluppo futuro del sistema."
    ), bold=True)

    doc.add_heading("10.2. Proposte di sviluppo e miglioramento", level=2)

    _p(doc, (
        "Interventi prioritari sui dati:"
    ), bold=True)

    _p(doc, (
        "• Rigeolocalizzazione sistematica di tutti gli incidenti del periodo 2010–2024, "
        "con aggancio diretto alla rete TomTom.\n"
        "• Introduzione della distinzione feriti gravi / feriti lievi nel flusso di "
        "registrazione (conforme alla definizione del Codice della Strada e alle "
        "convenzioni ISTAT).\n"
        "• Bonifica del campo «natura dell'incidente» e della classificazione per tipo "
        "di utente coinvolto (pedone, ciclista, motociclista).\n"
        "• Aggiornamento periodico della rete TomTom e dei dati di traffico/velocità."
    ))

    _p(doc, (
        "Miglioramenti metodologici:"
    ), bold=True)

    _p(doc, (
        "• Calibrazione di modelli SPF stratificati per periodo temporale (pre/post COVID, "
        "pre/post interventi infrastrutturali) per catturare trend e discontinuità.\n"
        "• Introduzione di modelli a effetti misti (random effects) per catturare "
        "l'eterogeneità spaziale non osservata.\n"
        "• Integrazione di covariate aggiuntive: sezione trasversale, poli attrattori "
        "(scuole, ospedali), illuminazione.\n"
        "• Analisi di sensitività strutturata con validazione rispetto a benchmark esterni "
        "(siti già noti come critici dalle Forze di Polizia).\n"
        "• Estensione dell'analisi ai ciclisti come categoria dedicata di utenti vulnerabili.\n"
        "• Modulo Before–After per la valutazione dell'efficacia degli interventi già realizzati.\n"
        "• Generazione automatica di schede-sito per i black point prioritari, con "
        "cartografia di dettaglio e suggerimenti di contromisure."
    ))

    doc.add_heading("10.3. Considerazioni finali", level=2)
    _p(doc, (
        "Il prototipo dimostra che l'approccio Empirical Bayes con indice composito "
        "multi-criterio è applicabile alla rete stradale di Roma e produce risultati "
        "coerenti con la letteratura e con l'esperienza sul campo. La pipeline è "
        "completamente automatizzata, riproducibile e aggiornabile. Tuttavia, la bontà "
        "dei risultati è intrinsecamente legata alla qualità dei dati incidentali: nessun "
        "modello statistico può compensare un dato di input sistematicamente incompleto "
        "o impreciso."
    ))

    _p(doc, (
        "Si raccomanda pertanto di subordinare l'adozione operativa del sistema a un "
        "investimento strutturale nella qualità dei dati incidentali, senza il quale "
        "le priorità individuate restano indicative ma non possono essere considerate "
        "definitive."
    ), bold=True)

    # ================================================================
    # 11. RIFERIMENTI BIBLIOGRAFICI
    # ================================================================
    doc.add_heading("11. Riferimenti bibliografici", level=1)

    refs = [
        "AASHTO (2010). Highway Safety Manual, 1st Edition. American Association of "
        "State Highway and Transportation Officials, Washington, D.C.",

        "Cameron, A.C. e Trivedi, P.K. (1998). Regression Analysis of Count Data. "
        "Cambridge University Press.",

        "Direttiva 2008/96/CE del Parlamento europeo e del Consiglio sulla gestione "
        "della sicurezza delle infrastrutture stradali. Gazzetta ufficiale dell'Unione "
        "europea, L 319/59.",

        "Direttiva (UE) 2019/1936 che modifica la direttiva 2008/96/CE. Gazzetta "
        "ufficiale dell'Unione europea, L 305/1.",

        "Hauer, E. (1997). Observational Before-After Studies in Road Safety. "
        "Pergamon Press, Oxford.",

        "Hauer, E., Harwood, D.W., Council, F.M. e Griffith, M.S. (2002). Estimating "
        "safety by the empirical Bayes method: a tutorial. Transportation Research "
        "Record, 1784(1), pp. 126–131.",

        "La Torre, F., Domenichini, L. e Corsi, F. (2019). A comparative analysis of "
        "network-level road safety screening methods. Accident Analysis & Prevention, "
        "130, pp. 272–281.",

        "Montella, A. (2010). A comparative analysis of hotspot identification methods. "
        "Accident Analysis & Prevention, 42(2), pp. 571–581.",

        "Persaud, B., Lyon, C. e Nguyen, T. (1999). Empirical Bayes procedure for "
        "ranking sites for safety investigation by potential for safety improvement. "
        "Transportation Research Record, 1665(1), pp. 7–12.",

        "MEF/ISTAT (2022). Valori monetari della sicurezza stradale – Costi sociali "
        "degli incidenti stradali. Nota metodologica, Roma.",

        "WHO (2018). Global Status Report on Road Safety. World Health Organization, "
        "Geneva.",

        "European Commission (2019). EU Road Safety Policy Framework 2021-2030 – "
        "Next steps towards \"Vision Zero\". Brussels.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)

    # ================================================================
    # SALVATAGGIO
    # ================================================================
    out = Path(__file__).resolve().parent / "relazione_metodologica.docx"
    doc.save(str(out))
    print(f"Relazione salvata: {out}")


if __name__ == "__main__":
    genera()
